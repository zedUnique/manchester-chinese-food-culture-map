from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT = Path("第四周总结.docx")
PAGE_IMAGE = Path("/private/tmp/week4-summary-page.png")
FONT_FILE = "/System/Library/Fonts/STHeiti Medium.ttc"

BLUE = "#2E74B5"
DARK_BLUE = "#1F4D78"
INK = "#1D2730"
MUTED = "#5B6870"
LIGHT_BLUE = "#E8EEF5"
LIGHT_GRAY = "#F4F6F9"
BORDER = "#C9D4E1"
WHITE = "#FFFFFF"

WIDTH = 975
HEIGHT = 1285
LEFT = 26
RIGHT = WIDTH - 26


def font(size):
    return ImageFont.truetype(FONT_FILE, size, index=0)


def wrap_text(text, text_font, max_width):
    lines = []
    current = ""
    for char in text:
        proposal = f"{current}{char}"
        if current and text_font.getlength(proposal) > max_width:
            lines.append(current)
            current = char
        else:
            current = proposal
    if current:
        lines.append(current)
    return lines


def draw_wrapped(draw, xy, text, text_font, fill, max_width, line_gap=7):
    x, y = xy
    bbox = text_font.getbbox("国")
    line_height = bbox[3] - bbox[1] + line_gap
    lines = wrap_text(text, text_font, max_width)
    for line in lines:
        draw.text((x, y), line, font=text_font, fill=fill)
        y += line_height
    return y


def draw_section_title(draw, y, text):
    text_font = font(25)
    draw.text((LEFT, y), text, font=text_font, fill=BLUE)
    return y + 38


def draw_bullet(draw, y, text):
    body_font = font(20)
    draw.ellipse((LEFT + 4, y + 10, LEFT + 11, y + 17), fill=BLUE)
    body_end = draw_wrapped(draw, (LEFT + 25, y), text, body_font, INK, RIGHT - LEFT - 25, line_gap=7)
    return body_end + 8


canvas = Image.new("RGB", (WIDTH, HEIGHT), WHITE)
draw = ImageDraw.Draw(canvas)

# Compact report header.
header_font = font(16)
header = "Manchester Chinese Food Culture Interactive Map | Week 4"
draw.text((WIDTH - 28 - header_font.getlength(header), 28), header, font=header_font, fill=MUTED)

kicker_font = font(19)
draw.text((LEFT, 80), "项目进度更新", font=kicker_font, fill=BLUE)

title_font = font(55)
draw.text((LEFT, 113), "第四周总结", font=title_font, fill=INK)

subtitle_font = font(27)
draw.text((LEFT, 183), "Manchester Chinese Food Culture Interactive Map", font=subtitle_font, fill=MUTED)
draw.line((LEFT, 225, RIGHT, 225), fill=BLUE, width=3)

metadata_font = font(17)
metadata = "学生: Bowen Zhang     阶段: 第四周     用途: 导师会议简报"
draw.text((LEFT, 243), metadata, font=metadata_font, fill=MUTED)

# Progress snapshot: genuinely tabular facts.
card_top, card_height, card_gap = 281, 79, 8
card_width = (RIGHT - LEFT - card_gap) // 2
cards = [
    ("27", "餐厅与地图点位"),
    ("57", "已整理菜品条目"),
    ("6", "饮食偏好/忌口筛选"),
    ("37", "菜品图片来源组"),
]
for index, (number, label) in enumerate(cards):
    row, column = divmod(index, 2)
    x1 = LEFT + column * (card_width + card_gap)
    y1 = card_top + row * (card_height + card_gap)
    x2, y2 = x1 + card_width, y1 + card_height
    draw.rounded_rectangle((x1, y1, x2, y2), radius=8, fill=LIGHT_GRAY, outline=BORDER, width=1)
    draw.text((x1 + 18, y1 + 16), number, font=font(33), fill=BLUE)
    draw.text((x1 + 78, y1 + 29), label, font=font(17), fill=INK)

y = 463
y = draw_section_title(draw, y, "1. 本周重点成果")
y = draw_bullet(draw, y, "完善互动地图：优化中国区域与曼城餐厅的联动逻辑，修正部分菜系归类、地址和标点信息，并改善餐厅选择后的跳转体验。")
y = draw_bullet(draw, y, "补充包容性信息：为菜品加入猪肉、牛肉、海鲜、麸质、花生/芝麻、大豆等提示，帮助有过敏、忌口或宗教饮食需求的用户筛选。")
y = draw_bullet(draw, y, "提升内容与视觉资料：更新区域饮食文化、菜品做法与故事；完成 57 个菜项的图片映射，并在图注中保留可追溯来源。")

y += 3
y = draw_section_title(draw, y, "2. 当前 Demo 已实现的使用流程")
flow_top = y
flow_bottom = flow_top + 76
draw.rounded_rectangle((LEFT, flow_top, RIGHT, flow_bottom), radius=8, fill=LIGHT_BLUE, outline=BORDER, width=1)
flow = "中国区域选择 → 区域饮食文化介绍 → 曼城餐厅地图定位 → 招牌菜选择 → 做法、口味、文化故事、图片与饮食标签展示"
draw_wrapped(draw, (LEFT + 16, flow_top + 13), flow, font(18), DARK_BLUE, RIGHT - LEFT - 32, line_gap=6)
y = flow_bottom + 20

y = draw_section_title(draw, y, "3. 下周优化方向")
y = draw_bullet(draw, y, "继续以餐厅官网、Google Maps 与实地观察核验地址、营业状态和代表菜，减少暂定信息。")
y = draw_bullet(draw, y, "替换部分通用参考图为更贴近曼城餐厅实际菜单的实拍素材，并继续保留可追溯来源。")
y = draw_bullet(draw, y, "进一步优化英文表达、移动端排版和地图信息层级，并邀请少量非中国用户进行可用性测试。")

takeaway_top = 1100
draw.rounded_rectangle((LEFT, takeaway_top, RIGHT, takeaway_top + 78), radius=8, fill="#F7FAFD", outline=BORDER, width=1)
draw.text((LEFT + 16, takeaway_top + 13), "当前阶段", font=font(18), fill=BLUE)
takeaway = "项目已从概念展示升级为可供导师反馈的互动原型；下一阶段聚焦资料核验、实拍补充与用户测试。"
draw_wrapped(draw, (LEFT + 16, takeaway_top + 37), takeaway, font(17), INK, RIGHT - LEFT - 32, line_gap=4)

footer_font = font(15)
footer = "Bowen Zhang | Fourth Week Summary"
draw.text(((WIDTH - footer_font.getlength(footer)) / 2, 1235), footer, font=footer_font, fill=MUTED)

canvas.save(PAGE_IMAGE, quality=95)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

doc.core_properties.title = "第四周总结"
doc.core_properties.author = "Bowen Zhang"
doc.core_properties.subject = "Manchester Chinese Food Culture Interactive Map - Week 4 Progress"
doc.core_properties.comments = "Meeting brief"

paragraph = doc.add_paragraph()
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph.paragraph_format.space_before = Pt(0)
paragraph.paragraph_format.space_after = Pt(0)
paragraph.paragraph_format.line_spacing = 1.0
run = paragraph.add_run()
run.add_picture(str(PAGE_IMAGE), width=Inches(6.5))

doc.save(OUTPUT)
print(f"Created {OUTPUT.resolve()}")
